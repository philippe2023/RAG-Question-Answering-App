# document_processing.py

import os
import logging
from typing import List
import hashlib

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import streamlit as st
import yaml

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

def load_config():
    with open("config.yaml", "r") as f:
        return yaml.safe_load(f)

config = load_config()

def is_document_already_processed(file_name: str) -> bool:
    """Checks if the document has already been processed."""
    # Implement your logic here, possibly by checking a database or a processed files list
    return False

def process_document(file, chunk_size: int = 1000, chunk_overlap: int = 100) -> List[Document]:
    """Processes an uploaded document, extracting text and splitting it into chunks."""
    try:
        file_extension = os.path.splitext(file.name)[1].lower()
        if file_extension == ".pdf":
            text = extract_text_from_pdf(file)
        elif file_extension == ".docx":
            text = extract_text_from_docx(file)
        elif file_extension == ".txt":
            text = extract_text_from_txt(file)
        elif file_extension == ".html":
            text = extract_text_from_html(file)
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return []

        # Split the text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=chunk_overlap
        )
        splits = text_splitter.split_text(text)
        # Create Document objects with metadata
        docs = []
        for idx, split in enumerate(splits):
            metadata = {
                "file_name": file.name,
                "chunk": idx,
            }
            doc = Document(page_content=split, metadata=metadata)
            docs.append(doc)
        return docs
    except Exception as e:
        logging.error(f"An error occurred while processing the document: {e}")
        st.error(f"An error occurred while processing the document: {e}")
        return []

def extract_text_from_pdf(file) -> str:
    """Extracts text from a PDF file."""
    try:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        logging.error(f"An error occurred while extracting text from PDF: {e}")
        st.error(f"An error occurred while extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(file) -> str:
    """Extracts text from a Word (.docx) file."""
    try:
        doc = DocxDocument(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logging.error(f"An error occurred while extracting text from DOCX: {e}")
        st.error(f"An error occurred while extracting text from DOCX: {e}")
        return ""

def extract_text_from_txt(file) -> str:
    """Extracts text from a text (.txt) file."""
    try:
        text = file.read().decode('utf-8')
        return text
    except Exception as e:
        logging.error(f"An error occurred while extracting text from TXT: {e}")
        st.error(f"An error occurred while extracting text from TXT: {e}")
        return ""

def extract_text_from_html(file) -> str:
    """Extracts text from an HTML file."""
    try:
        content = file.read().decode('utf-8')
        soup = BeautifulSoup(content, 'html.parser')
        text = soup.get_text(separator="\n")
        return text
    except Exception as e:
        logging.error(f"An error occurred while extracting text from HTML: {e}")
        st.error(f"An error occurred while extracting text from HTML: {e}")
        return ""